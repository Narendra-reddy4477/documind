import re
from pathlib import Path
from typing import List
from langchain.schema import Document


def parse_pdf(file_path: str) -> List[Document]:
    """Extract text from PDF with page tracking."""
    from pypdf import PdfReader
    reader = PdfReader(file_path)
    documents = []
    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = _clean_text(text)
        if text.strip():
            documents.append(Document(
                page_content=text,
                metadata={"page": page_num, "source": Path(file_path).name}
            ))
    return documents


def parse_docx(file_path: str) -> List[Document]:
    """Extract text from DOCX preserving paragraph structure."""
    from docx import Document as DocxDoc
    doc = DocxDoc(file_path)
    documents = []
    current_text = []
    page_estimate = 1

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        current_text.append(text)
        # Rough page estimate every ~40 paragraphs
        if len(current_text) % 40 == 0:
            page_estimate += 1

    full_text = "\n".join(current_text)
    full_text = _clean_text(full_text)

    if full_text.strip():
        documents.append(Document(
            page_content=full_text,
            metadata={"page": 1, "source": Path(file_path).name}
        ))
    return documents


def parse_txt(file_path: str) -> List[Document]:
    """Extract text from plain text files."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    text = _clean_text(text)
    if text.strip():
        return [Document(
            page_content=text,
            metadata={"page": 1, "source": Path(file_path).name}
        )]
    return []


def parse_document(file_path: str, file_type: str) -> List[Document]:
    """Route to correct parser based on file type."""
    parsers = {
        "pdf":  parse_pdf,
        "docx": parse_docx,
        "txt":  parse_txt,
    }
    parser = parsers.get(file_type.lower())
    if not parser:
        raise ValueError(f"Unsupported file type: {file_type}")
    return parser(file_path)


def _clean_text(text: str) -> str:
    """Remove excessive whitespace and non-printable characters."""
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\x20-\x7E\n]', '', text)
    return text.strip()
